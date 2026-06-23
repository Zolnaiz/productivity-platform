from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
OUT_FILE = OUT_DIR / "Productivity_Platform_Project_Report_2026-06-23.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, width in (("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), width)
                node.set(qn("w:type"), "dxa")

    document.add_paragraph()


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    run = paragraph.runs[0]
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 12)
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 12 if level == 2 else 8)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)


def add_para(document: Document, text: str, bold_label: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_label and text.startswith(bold_label):
        first = paragraph.add_run(bold_label)
        first.bold = True
        first.font.name = "Arial"
        first.font.size = Pt(11)
        rest = paragraph.add_run(text[len(bold_label) :])
        rest.font.name = "Arial"
        rest.font.size = Pt(11)
    else:
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Productivity Platform төслийн дэлгэрэнгүй тайлан")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Зорилго, одоогийн төлөв, хийсэн ажил, цаашдын roadmap")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(90, 90, 90)

    add_table(
        document,
        ["Талбар", "Мэдээлэл"],
        [
            ["Огноо", "2026-06-23"],
            ["Зориулалт", "Төслийн эзэн, хөгжүүлэлтийн баг, ирээдүйн дизайн/логик шийдвэр гаргалт"],
            ["Хамрах хүрээ", "Web admin, backend API, mobile Flutter app, Figma UI blueprint, 5S/assessment/reporting roadmap"],
            ["Figma дизайн", "https://www.figma.com/design/mqU0VidETu0ipGvfGQBG9d"],
            ["FigJam зураглал", "https://www.figma.com/board/mczFd8ZmcxXWmpmj36nICE"],
        ],
    )

    add_heading(document, "1. Товч дүгнэлт", 1)
    add_para(
        document,
        "Энэ төсөл нь байгууллага, баг, ажилтан бүрийн төлөвлөлт, өдөр тутмын гүйцэтгэл, цаг ашиглалт, төсөл, тайлан, 5S/audit/checklist, үнэлгээ, зардал болон productivity tracking-ийг нэг дор удирдах платформ болох зорилготой.",
    )
    add_para(
        document,
        "Одоогийн байдлаар backend/API, web admin, mobile Flutter суурь ажиллах хэмжээнд тогтворжсон. Web дээр assessment/checklist хэсгийг өргөжүүлж, mobile APK/emulator ажиллуулах асуудлыг засаж, Figma дээр бүрэн бүтээгдэхүүний UI blueprint болон prototype урсгалыг гаргасан.",
    )
    add_para(
        document,
        "Гол дараагийн ажил бол Figma дээр гаргасан бүрэн зураглалыг code руу үе шаттай оруулах: 5S implementation setup, 5S audit run, role/permission, report analytics, mobile 5S flow, light/dark theme, production deployment hardening.",
    )

    add_heading(document, "2. Бүтээгдэхүүний зорьж буй бүрэн дүр зураг", 1)
    add_bullets(
        document,
        [
            "Owner/Admin: байгууллага, хэрэглэгч, эрх, стандарт, тайлан, аудитын бүх тохиргоог удирдана.",
            "Manager: төсөл, ажил, багийн гүйцэтгэл, checklist, 5S хэрэгжилт, тайланг хянана.",
            "Employee: өөрийн ажил, цаг, өдөр тутмын log, checklist, audit evidence, expense request-ээ бүртгэнэ.",
            "Auditor/Quality role: 5S, safety, quality, compliance audit хийж, оноо, evidence, corrective action гаргана.",
            "Салбарын өргөтгөл: hospitality, manufacturing, construction, retail, logistics, facility management-д template-ээр тохируулна.",
        ],
    )

    add_table(
        document,
        ["Давхарга", "Зорилго", "Хэрэглэгчид харагдах үнэ цэнэ"],
        [
            ["Backend/API", "Өгөгдөл, auth, role, report, audit, task workflow-г нэг API-д төвлөрүүлэх", "Нэг эх сурвалжтай, тайлан болон mobile/web ижил мэдээлэл ашиглана"],
            ["Admin Web", "Менежер/admin өдөр тутмын удирдлага хийх үндсэн interface", "Dashboard, projects, tasks, assessment, reports, people, settings"],
            ["Mobile App", "Ажилтан талын хурдан бүртгэл, field audit, evidence capture", "Гар утсаар log, checklist, expense, profile, 5S audit хийх"],
            ["Figma/Product Design", "Бүрэн бүтээгдэхүүний UI, мэдээллийн бүтэц, prototype-г урьдчилан харах", "Юу нэмэх, яаж ажиллахыг code бичихээс өмнө ойлгомжтой болгоно"],
        ],
    )

    add_heading(document, "3. Төлөвлөсөн модуль ба одоогийн төлөв", 1)
    add_table(
        document,
        ["Модуль", "Гол агуулга", "Одоогийн төлөв"],
        [
            ["Auth / Organization", "Нэвтрэх, бүртгүүлэх, JWT, organization scope, role guard", "Суурь хийгдсэн, hardening хийгдсэн"],
            ["Projects / Tasks / Kanban", "Төсөл, ажил, хариуцагч, хугацаа, priority, progress", "Web/API суурь хийгдсэн, UX сайжруулах шаардлагатай"],
            ["Work logs / Time tracking", "Өдөр тутмын хийсэн ажил, цаг, blocker, next step", "Суурь төлөвлөгөө бэлэн, гүнзгий workflow нэмэх шаардлагатай"],
            ["Assessment / Checklist", "Template builder, question types, response, score, review", "Web дээр өргөжүүлэлт хийгдсэн"],
            ["5S Implementation", "Байгууллага/site/floor/zone/asset/standard/owner/schedule setup", "Figma дизайн бэлэн, code хэрэгжилт дараагийн шат"],
            ["5S Audit Run", "Sort, Set in order, Shine, Standardize, Sustain score, evidence, corrective action", "Figma дизайн бэлэн, code хэрэгжилт дараагийн шат"],
            ["Reports / Analytics", "Сар, ажилтан, төсөл, department, export", "Monthly report суурь болон service coverage бий, dashboard-тай нягт холбох хэрэгтэй"],
            ["Expenses", "Зардлын хүсэлт, approval, report", "Mobile/web суурь байгаа, approval workflow гүнзгийрүүлэх хэрэгтэй"],
            ["People / Roles", "User, department, permission matrix", "Role guard суурь бий, permission matrix дизайн бэлэн"],
            ["Theme / Design System", "Light/dark, component tokens, responsive UI", "Figma blueprint бэлэн, code руу бүрэн шилжүүлэх хэрэгтэй"],
        ],
    )

    add_heading(document, "4. Одоо хийгдсэн гол ажлууд", 1)
    add_heading(document, "4.1 Backend/API", 2)
    add_bullets(
        document,
        [
            "Auth болон organization migration, demo organization/owner seed, JWT login/refresh/profile contract тогтворжсон.",
            "Project CRUD smoke, soft delete, organization scope guard, DTO validation, nested checklist/answer validation нэмэгдсэн.",
            "Monthly report filtering, metrics endpoint, structured logging, request ID, rate limiting, exception filter hardening хийгдсэн.",
            "Docker compose production/dev/test тохиргоо, production secrets validation, backup/restore guidance бэлэн болсон.",
        ],
    )

    add_heading(document, "4.2 Web admin", 2)
    add_bullets(
        document,
        [
            "Dashboard, projects, tasks/kanban, work logs, reports, people, finance, settings зэрэг үндсэн page-үүд суурьтай болсон.",
            "Questionnaires хэсгийг Assessments & Checklists болгон өргөжүүлж, multi-question builder, quick-start preset, publish/archive workflow нэмсэн.",
            "Responses хэсэгт template сонгох, dynamic answer бөглөх, score тооцох, review/reject, low score improvement task үүсгэх урсгал нэмэгдсэн.",
            "Route шинэчлэлтээр `/assessments` ажиллаж, хуучин `/questionnaires` redirect хийгддэг болсон.",
            "`npm run build` амжилттай давсан, `/assessments` болон `/responses` web route HTTP 200 буцаасан.",
        ],
    )

    add_heading(document, "4.3 Mobile Flutter", 2)
    add_bullets(
        document,
        [
            "Android SDK/NDK/Gradle тохиргоо засагдаж APK build/install ажилласан.",
            "APK байршил: C:\\Users\\NewTech\\Desktop\\productivity-platform\\mobile-flutter\\build\\app\\outputs\\flutter-apk\\app-dev-debug.apk",
            "GoRouter болон bottom navigation засагдаж Dashboard, Forms/Assessments, Expenses, Reports, Profile хооронд шилжилт тогтвортой болсон.",
            "Profile overflow болон UI-ийн зарим layout асуудлыг зассан, logcat дээр fatal/navigation/overflow алдаа илрээгүй.",
        ],
    )

    add_heading(document, "4.4 Figma дизайн", 2)
    add_bullets(
        document,
        [
            "Bilguun's team дотор complete UI blueprint үүсгэсэн.",
            "Dark/light mode, web dashboard, projects/tasks, assessment builder, reports, people/settings, mobile screens, 5S implementation, 5S audit, role matrix, data/API model зэрэг frame-үүд бэлэн.",
            "Prototype flow map нэмэгдсэн: home -> dashboard -> projects -> 5S setup -> audit run -> assessments -> reports -> settings.",
            "5S implementation болон 5S audit-ийг тусдаа логиктой болгож дизайн дээр салгасан.",
        ],
    )

    add_heading(document, "5. 5S-ийн зөв бүтээгдэхүүний логик", 1)
    add_para(document, "5S Implementation болон 5S Audit нь нэг зүйл биш. Implementation нь байгууллага 5S-ийг яаж хэрэгжүүлэх бүтэц, стандарт, эзэмшил, давтамжийг тохируулах хэсэг. Audit нь тэр тохиргооны дагуу бодит орчинд оноо, evidence, finding, corrective action бүртгэх хэсэг.")
    add_table(
        document,
        ["5S Implementation setup", "Жишээ өгөгдөл"],
        [
            ["Organization / Site", "Компани, салбар, үйлдвэр, оффис"],
            ["Floor / Department / Zone", "2-р давхар, агуулах, гал тогоо, reception, workshop"],
            ["Asset / Area item", "Ширээ, шүүгээ, багажны самбар, rack, emergency exit"],
            ["Standard", "Юу байх ёстой, зураг/жишиг, label, cleaning rule, max/min stock"],
            ["Owner", "Хариуцагч ажилтан, supervisor, backup owner"],
            ["Schedule", "Өдөр бүр, 7 хоног бүр, сар бүр audit хийх давтамж"],
        ],
    )
    add_table(
        document,
        ["5S Audit run", "Жишээ өгөгдөл"],
        [
            ["Sort", "Хэрэггүй зүйлс байна уу, устгах/шилжүүлэх шаардлагатай юу"],
            ["Set in order", "Байршил, label, access, visual control зөв үү"],
            ["Shine", "Цэвэрлэгээ, эвдрэл, аюулгүй байдал хангагдсан уу"],
            ["Standardize", "Стандарт, SOP, зураг, checklist баримтлагдаж байна уу"],
            ["Sustain", "Давтамж, ownership, training, previous action follow-up тогтвортой юу"],
            ["Evidence / Corrective action", "Зураг, тэмдэглэл, root cause, due date, responsible person"],
        ],
    )

    add_heading(document, "6. Одоогийн техникийн төлөв ба анхаарах зүйл", 1)
    add_bullets(
        document,
        [
            "Web admin local орчинд `http://127.0.0.1:3001` дээр ажиллаж байсан.",
            "Backend API local орчинд `http://127.0.0.1:3000` чиглэлээр хэрэглэгдэж байна.",
            "Docker зарим үед орчны/daemon асуудлаас болж төвөгтэй байж болно; backend smoke нь өмнө PostgreSQL-той давсан гэж roadmap-д тэмдэглэгдсэн.",
            "Production-д `JWT_SECRET`, database credentials, CORS origins зэрэг real secret заавал тусдаа тохируулах шаардлагатай.",
            "Sensitive production data орох бол localStorage token storage-ийг HttpOnly cookie/session стратеги руу шилжүүлэх нь зөв.",
            "Figma дээрх бүрэн дизайн code руу хараахан бүрэн шилжээгүй, энэ нь дараагийн хамгийн том ажил.",
        ],
    )

    add_heading(document, "7. Цаашдын roadmap", 1)
    add_table(
        document,
        ["Үе шат", "Хийх ажил", "Гарах үр дүн"],
        [
            ["Phase 1", "Figma blueprint-ийг web navigation, dashboard, theme, responsive layout-д оруулах", "Одоогийн bolhi UI багасаж, ойлгомжтой modern web experience болно"],
            ["Phase 2", "5S Implementation backend entity/API/web UI нэмэх", "Байгууллагын бүтэц, area, standard, owner, schedule бүртгэдэг болно"],
            ["Phase 3", "5S Audit Run, evidence, finding, corrective action, score workflow нэмэх", "Бодит audit хийж, score болон action plan гардаг болно"],
            ["Phase 4", "Assessment/Checklist-ийг report analytics, improvement task, notification-той холбох", "Доогуур онооны асуудлууд автоматаар ажил болж хувирна"],
            ["Phase 5", "Work log/time tracking/monthly report export-ийг бүрэн болгох", "Ажилтан, баг, төсөл бүрийн сарын тайлан автоматаар гарна"],
            ["Phase 6", "Mobile 5S flow, offline-friendly capture, camera evidence, sync", "Field ажилтан утсаар audit/evidence оруулах боломжтой болно"],
            ["Phase 7", "Role permission matrix, approval workflow, production deploy/monitoring", "Бодит байгууллагад ашиглахад бэлэн production шатанд очно"],
        ],
    )

    add_heading(document, "8. Шийдэх шаардлагатай нээлттэй асуудлууд", 1)
    add_bullets(
        document,
        [
            "UI хэл: Монгол/Англи эсвэл i18n олон хэлний бүтэцтэй явах эсэх.",
            "5S hierarchy яг ямар нэршилтэй байх: site/floor/department/zone/asset гэсэн universal model эсвэл салбар бүрээр тусдаа model ашиглах эсэх.",
            "Audit evidence зураг/файлыг local, S3-compatible storage, эсвэл cloud provider дээр хадгалах эсэх.",
            "Production deployment target: local server, VPS, Render, Netlify + backend host, Cloudflare зэрэг сонголтуудаас аль нь тохирох.",
            "Role permission нарийвчлал: owner/admin/manager/employee/auditor/finance гэх мэт role-уудын эрхийг эцэслэх.",
        ],
    )

    add_heading(document, "9. Зөвлөмж", 1)
    add_para(
        document,
        "Ойрын ажлыг 5S Implementation setup-ээс эхлэх нь хамгийн зөв. Учир нь audit хийхээс өмнө байгууллагын area, standard, owner, frequency, target score тодорхой байх ёстой. Дараа нь Audit Run, Corrective Action, Reports, Mobile Evidence Capture гэсэн дарааллаар нэмбэл бүтээгдэхүүний логик цэвэрхэн өснө.",
    )
    add_para(
        document,
        "Figma дээр гаргасан complete UI blueprint-ийг product truth болгон ашиглаж, web/mobile code-ыг тэр дизайнтай үе шаттай тааруулж явбал дараа нь module нэмэхэд замбараатай болно.",
    )

    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    add_heading(document, "Хавсралт: Одоогийн өөрчлөлтийн тойм", 1)
    add_table(
        document,
        ["Төрөл", "Файл/хэсэг", "Тайлбар"],
        [
            ["Web", "admin-web/src/App.tsx", "Assessments route нэмэгдэж, legacy questionnaire route redirect хийгдсэн"],
            ["Web", "admin-web/src/components/layout/Sidebar.tsx", "Questionnaires нэршил Assessments болж navigation шинэчлэгдсэн"],
            ["Web", "admin-web/src/pages/QuestionnairesPage.tsx", "Assessment/checklist builder болгож өргөжүүлсэн"],
            ["Web", "admin-web/src/pages/ResponsesPage.tsx", "Response бөглөх, score, review, improvement task урсгал нэмэгдсэн"],
            ["Mobile", "mobile-flutter/lib/app_router.dart", "Navigation/button шилжилтийн асуудал засагдсан"],
            ["Mobile", "mobile-flutter/lib/screens/*.dart", "Dashboard, questionnaire, expense, profile зэрэг screen-ийн тогтвортой байдал сайжирсан"],
            ["Android", "mobile-flutter/android/*", "Gradle/SDK/APK/emulator build тохиргоо засагдсан"],
            ["Figma", "Productivity Platform Complete UI Blueprint", "Бүрэн UI, prototype, 5S logic, roadmap frame-үүд үүссэн"],
        ],
    )

    document.save(OUT_FILE)


if __name__ == "__main__":
    build_document()
    print(OUT_FILE)
